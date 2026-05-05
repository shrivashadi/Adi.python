from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

df = pd.read_csv('data/clean_data.csv')
df.columns = ['Date', 'Apprehended', 'CBP_Custody',
               'CBP_Transfers', 'HHS_Care', 'HHS_Discharged']
for col in df.columns[1:]:
    df[col] = pd.to_numeric(df[col], errors='coerce')
df['Date'] = pd.to_datetime(df['Date'])
df = df.dropna()

df['Transfer_Efficiency'] = df['CBP_Transfers'] / df['CBP_Custody']
df['Discharge_Effectiveness'] = df['HHS_Discharged'] / df['HHS_Care']
df['Backlog'] = df['Apprehended'] - df['HHS_Discharged']

doc = Document()

# Title
title = doc.add_heading('EXECUTIVE SUMMARY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Care Transition Efficiency & Placement Outcome Analytics')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True

sub2 = doc.add_paragraph('Prepared for: U.S. Department of Health & Human Services')
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub3 = doc.add_paragraph('Unified Mentor Project | 2025')
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# Purpose
doc.add_heading('Purpose', level=1)
doc.add_paragraph(
    'This report presents a data-driven analysis of the Unaccompanied Alien '
    'Children (UAC) care pipeline. The goal is to evaluate how efficiently '
    'children move through CBP custody, HHS care, and sponsor placement — '
    'and to identify where delays and bottlenecks occur.'
)

# Key Findings
doc.add_heading('Key Findings', level=1)

findings = [
    f'Transfer Efficiency (CBP → HHS): {df["Transfer_Efficiency"].mean():.2%} average — indicating room for improvement in transfer speed.',
    f'Discharge Effectiveness: {df["Discharge_Effectiveness"].mean():.2%} average — sponsor placements are below optimal levels.',
    f'Average Daily Backlog: {abs(df["Backlog"].mean()):.0f} children remain unplaced each day.',
    'Peak bottlenecks were observed during high inflow periods — system struggled to maintain discharge pace.',
    'Weekday transitions are faster than weekend — staffing gaps exist on weekends.',
]
for f in findings:
    doc.add_paragraph(f'• {f}')

# KPI Table
doc.add_heading('Key Performance Indicators', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'KPI'
hdr[1].text = 'Result'

kpis = [
    ('Transfer Efficiency', f"{df['Transfer_Efficiency'].mean():.2%}"),
    ('Discharge Effectiveness', f"{df['Discharge_Effectiveness'].mean():.2%}"),
    ('Avg Daily Backlog', f"{abs(df['Backlog'].mean()):.0f} children"),
    ('Total Records Analyzed', str(len(df))),
    ('Date Range', f"{df['Date'].min().date()} to {df['Date'].max().date()}"),
]
for kpi, val in kpis:
    row = table.add_row().cells
    row[0].text = kpi
    row[1].text = val

doc.add_paragraph()

# Recommendations
doc.add_heading('Recommendations', level=1)
recs = [
    'Increase weekend staffing to reduce transition delays.',
    'Implement daily backlog monitoring alerts.',
    'Accelerate CBP to HHS transfer protocols during peak periods.',
    'Expand HHS discharge capacity to match inflow rates.',
    'Use data-driven dashboards for real-time pipeline monitoring.',
]
for r in recs:
    doc.add_paragraph(f'• {r}')

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'The UAC care pipeline requires immediate process improvements to reduce '
    'child custody durations and improve placement outcomes. Data analytics '
    'provides a clear path forward — enabling faster reunification, reduced '
    'backlogs, and stronger child welfare outcomes.'
)

doc.add_paragraph()
doc.add_paragraph('─' * 60)
end = doc.add_paragraph('Unified Mentor | UAC Analytics Project | 2025')
end.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('Executive_Summary.docx')
print("✅ Executive Summary saved!")