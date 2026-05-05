from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# Load Data
df = pd.read_csv('data/clean_data.csv')
df.columns = ['Date', 'Apprehended', 'CBP_Custody',
               'CBP_Transfers', 'HHS_Care', 'HHS_Discharged']
for col in df.columns[1:]:
    df[col] = pd.to_numeric(df[col], errors='coerce')
df['Date'] = pd.to_datetime(df['Date'])
df = df.dropna()

df['Transfer_Efficiency'] = df['CBP_Transfers'] / df['CBP_Custody']
df['Discharge_Effectiveness'] = df['HHS_Discharged'] / df['HHS_Care']
df['Backlog'] = df['Apprehended'] - df['HHS_Discharged']
df['Pipeline_Throughput'] = (df['CBP_Transfers'] + df['HHS_Discharged']) / df['Apprehended']

doc = Document()

# Title
title = doc.add_heading('Care Transition Efficiency & Placement Outcome Analytics', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph('Research Report | Unified Mentor Project | 2025')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Abstract
doc.add_heading('1. Abstract', level=1)
doc.add_paragraph(
    'This research analyzes the U.S. Unaccompanied Alien Children (UAC) program '
    'care pipeline efficiency. Using data from the HHS program, we evaluate key '
    'performance indicators including Transfer Efficiency, Discharge Effectiveness, '
    'and Pipeline Throughput to identify bottlenecks and recommend improvements.'
)

# Introduction
doc.add_heading('2. Introduction & Background', level=1)
doc.add_paragraph(
    'The UAC Program operates as a multi-stage care and reunification pipeline. '
    'The pipeline consists of: Apprehension & CBP custody, Transfer to HHS care, '
    'Medical screening, sheltering, and case management, and Discharge and '
    'reunification with a vetted sponsor. From a policy and humanitarian perspective, '
    'speed, continuity, and reliability of this pipeline are as critical as capacity.'
)

# Problem Statement
doc.add_heading('3. Problem Statement', level=1)
doc.add_paragraph(
    'While aggregate counts of children in custody are monitored, process efficiency '
    'metrics are largely absent. Key unanswered questions include: How efficiently '
    'are children transferred from CBP to HHS? Are discharges keeping pace with '
    'inflows? When and where do care backlogs accumulate? Are placement outcomes '
    'improving or deteriorating over time? Without structured transition analytics, '
    'system bottlenecks remain hidden.'
)

# Dataset Description
doc.add_heading('4. Dataset Description', level=1)
doc.add_paragraph(f'Total Records: {len(df)}')
doc.add_paragraph(f'Date Range: {df["Date"].min().date()} to {df["Date"].max().date()}')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Column'
hdr[1].text = 'Description'

columns_desc = [
    ('Date', 'Reporting date'),
    ('Apprehended', 'Daily intake volume'),
    ('CBP_Custody', 'Active CBP care load'),
    ('CBP_Transfers', 'Flow into HHS system'),
    ('HHS_Care', 'Active HHS care load'),
    ('HHS_Discharged', 'Successful sponsor placements'),
]
for col, desc in columns_desc:
    row = table.add_row().cells
    row[0].text = col
    row[1].text = desc

doc.add_paragraph()

# Methodology
doc.add_heading('5. Analytical Methodology', level=1)
doc.add_paragraph('The following KPIs were derived:')
doc.add_paragraph('• Transfer Efficiency = CBP Transfers ÷ CBP Custody', style='List Bullet')
doc.add_paragraph('• Discharge Effectiveness = HHS Discharged ÷ HHS Care', style='List Bullet')
doc.add_paragraph('• Backlog = Apprehended - HHS Discharged', style='List Bullet')
doc.add_paragraph('• Pipeline Throughput = (Transfers + Discharged) ÷ Apprehended', style='List Bullet')

# KPI Results
doc.add_heading('6. KPI Results', level=1)

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'KPI'
hdr2[1].text = 'Value'

kpis = [
    ('Avg Transfer Efficiency', f"{df['Transfer_Efficiency'].mean():.2%}"),
    ('Avg Discharge Effectiveness', f"{df['Discharge_Effectiveness'].mean():.2%}"),
    ('Avg Daily Backlog', f"{df['Backlog'].mean():.0f} children"),
    ('Avg Pipeline Throughput', f"{df['Pipeline_Throughput'].mean():.2%}"),
]
for kpi, val in kpis:
    row = table2.add_row().cells
    row[0].text = kpi
    row[1].text = val

doc.add_paragraph()

# Add Charts
doc.add_heading('7. EDA Findings', level=1)
try:
    doc.add_picture('pipeline_flow.png', width=Inches(6))
    doc.add_paragraph('Figure 1: UAC Care Pipeline Flow Over Time').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture('backlog.png', width=Inches(6))
    doc.add_paragraph('Figure 2: Daily Backlog Accumulation').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture('monthly_trends.png', width=Inches(6))
    doc.add_paragraph('Figure 3: Monthly KPI Trends').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture('correlation.png', width=Inches(6))
    doc.add_paragraph('Figure 4: Correlation Heatmap').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('Note: Run Jupyter Notebook first to generate charts.')

# Recommendations
doc.add_heading('8. Recommendations', level=1)
doc.add_paragraph('• Improve CBP to HHS transfer speed to reduce custody duration.', style='List Bullet')
doc.add_paragraph('• Increase HHS discharge capacity during high inflow periods.', style='List Bullet')
doc.add_paragraph('• Monitor backlog accumulation daily to detect bottlenecks early.', style='List Bullet')
doc.add_paragraph('• Implement weekend staffing to maintain transfer efficiency.', style='List Bullet')

# Conclusion
doc.add_heading('9. Conclusion', level=1)
doc.add_paragraph(
    'This project reframes the UAC dataset from a capacity monitoring lens to a '
    'process efficiency and outcome evaluation lens. By analyzing how effectively '
    'children move through the care pipeline, it provides actionable insights for '
    'improving reunification timelines, reducing delays, and strengthening child '
    'welfare outcomes.'
)

# Save
doc.save('UAC_Research_Paper.docx')
print("✅ Research Paper saved as UAC_Research_Paper.docx")